"""Phase 2: 文档处理管道 — 解析 → 分块 → Embedding → 存储"""
import os
import logging
from flask import current_app

from app.extensions import db
from app.models.knowledge import KnowledgeCollection, KnowledgeDocument
from app.core.knowledge.embedder import EmbeddingService

logger = logging.getLogger(__name__)

SUPPORTED_TYPES = {
    "pdf": "pypdf",
    "docx": "docx",
    "md": "text",
    "txt": "text",
    "html": "text",
    "csv": "text",
}


class DocumentPipeline:
    def __init__(self, embedding_service: EmbeddingService = None):
        self.embedder = embedding_service or EmbeddingService()

    def process(self, file_data: bytes, filename: str, collection_id: int,
                tenant_id: int, storage_dir: str) -> KnowledgeDocument:
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "txt"
        if ext not in SUPPORTED_TYPES:
            raise ValueError(f"不支持的文件类型: .{ext}，支持: {list(SUPPORTED_TYPES.keys())}")

        # 1. 创建文档记录
        doc = KnowledgeDocument(
            collection_id=collection_id,
            tenant_id=tenant_id,
            filename=filename,
            file_type=ext,
            file_size=len(file_data),
            status="parsing",
        )
        db.session.add(doc)
        db.session.commit()

        try:
            # 2. 保存原始文件
            os.makedirs(storage_dir, exist_ok=True)
            file_path = os.path.join(storage_dir, f"{doc.id}_{filename}")
            with open(file_path, "wb") as f:
                f.write(file_data)
            doc.file_path = file_path
            db.session.commit()

            # 3. 解析文本
            raw_text = self._parse(file_data, ext)

            # 4. 分块
            chunks = self._split(raw_text)

            if not chunks:
                raise ValueError("文档解析后无有效文本内容")

            # 5. 生成 Embedding 并存储到 pgvector
            self._store_chunks(chunks, doc, collection_id, tenant_id)

            # 6. 更新状态
            doc.status = "ready"
            doc.chunk_count = len(chunks)
            db.session.commit()
            logger.info(f"Document {doc.id} processed: {len(chunks)} chunks")

        except Exception as e:
            doc.status = "error"
            doc.error_message = str(e)
            db.session.commit()
            logger.exception(f"Document {doc.id} processing failed")
            raise

        return doc

    def _parse(self, data: bytes, ext: str) -> str:
        if ext == "pdf":
            from io import BytesIO
            from pypdf import PdfReader
            reader = PdfReader(BytesIO(data))
            return "\n\n".join(
                page.extract_text() or "" for page in reader.pages
            )
        elif ext == "docx":
            from io import BytesIO
            from docx import Document
            doc = Document(BytesIO(data))
            return "\n\n".join(p.text for p in doc.paragraphs)
        else:
            return data.decode("utf-8", errors="replace")

    def _split(self, text: str) -> list[str]:
        from langchain_text_splitters import RecursiveCharacterTextSplitter
        collection = None
        # Try to get collection config
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            separators=["\n\n", "\n", "。", "！", "？", ".", "!", "?", " ", ""],
        )
        return [d.page_content for d in splitter.create_documents([text])]

    def _store_chunks(self, chunks: list[str], doc: KnowledgeDocument,
                      collection_id: int, tenant_id: int):
        if not chunks:
            return
        embeddings = self.embedder.embed_documents(chunks)
        table_name = self._table_name(collection_id)
        dim = len(embeddings[0])
        self._ensure_table(table_name, dim)

        import json as _json
        # 通过原始 psycopg2 连接插入，绕过 SQLAlchemy text() 的类型问题
        raw_conn = db.engine.raw_connection()
        cur = raw_conn.cursor()
        try:
            for i, (chunk, emb) in enumerate(zip(chunks, embeddings)):
                emb_str = "[" + ",".join(str(x) for x in emb) + "]"
                meta = _json.dumps({
                    "document_id": doc.id, "collection_id": collection_id,
                    "tenant_id": tenant_id, "filename": doc.filename,
                }, ensure_ascii=False)
                cur.execute(
                    f"INSERT INTO {table_name} (document_id, chunk_index, content, embedding, metadata) "
                    f"VALUES (%s, %s, %s, %s::vector, %s::jsonb)",
                    (doc.id, i, chunk, emb_str, meta)
                )
            raw_conn.commit()
        finally:
            cur.close()
            raw_conn.close()

    def _table_name(self, collection_id: int) -> str:
        return f"kb_chunks_{collection_id}"

    def _ensure_table(self, table_name: str, dim: int):
        from sqlalchemy import text, inspect
        if not inspect(db.engine).has_table(table_name):
            db.session.execute(text(
                f"CREATE TABLE {table_name} ("
                f"  id SERIAL PRIMARY KEY,"
                f"  document_id INTEGER,"
                f"  chunk_index INTEGER,"
                f"  content TEXT,"
                f"  embedding vector({dim}),"
                f"  metadata JSONB"
                f")"
            ))
            db.session.execute(text(
                f"CREATE INDEX IF NOT EXISTS idx_{table_name}_doc ON {table_name}(document_id)"
            ))
            db.session.commit()
            logger.info(f"Created pgvector table: {table_name} (dim={dim})")

    def delete_document_chunks(self, document_id: int, collection_id: int):
        from sqlalchemy import text
        table = self._table_name(collection_id)
        db.session.execute(text(f"DELETE FROM {table} WHERE document_id = :did"), {"did": document_id})
        db.session.commit()
